"""Document templates: Contract generation, Letters, CV, PDF rendering."""

from dataclasses import dataclass
from pathlib import Path
from typing import Any, ClassVar, cast

from jinja2 import Environment, FileSystemLoader


@dataclass
class TemplateContext:
    variables: dict[str, Any]
    locale: str = "en"


class TemplateEngine:
    def __init__(self, template_dir: str = "templates") -> None:
        self._env = Environment(
            loader=FileSystemLoader(template_dir),
            autoescape=True,
            trim_blocks=True,
            lstrip_blocks=True,
        )
        self._add_filters()

    def _add_filters(self) -> None:
        self._env.filters["currency"] = lambda v, c="EUR": f"{c} {v:,.2f}"
        self._env.filters["date"] = lambda v, f="%Y-%m-%d": v.strftime(f) if v else ""
        self._env.filters["upper"] = str.upper
        self._env.filters["lower"] = str.lower
        self._env.filters["nl2br"] = lambda v: v.replace("\n", "<br>") if v else ""

    def render(self, template_name: str, context: dict[str, Any]) -> str:
        template = self._env.get_template(template_name)
        return template.render(**context)

    def render_to_file(self, template_name: str, context: dict[str, Any], output_path: str) -> None:
        content = self.render(template_name, context)
        Path(output_path).write_text(content)


class PDFRenderer:
    def __init__(self) -> None:
        pass

    def html_to_pdf(self, html: str, output: str | None = None) -> bytes:
        from weasyprint import HTML

        doc = HTML(string=html)
        pdf_bytes = cast("bytes", doc.write_pdf())

        if output:
            with open(output, "wb") as f:
                f.write(pdf_bytes)

        return pdf_bytes

    def render_template_to_pdf(
        self,
        template_name: str,
        context: dict[str, Any],
        template_dir: str,
        output: str | None = None,
    ) -> bytes:
        engine = TemplateEngine(template_dir)
        html = engine.render(template_name, context)
        return self.html_to_pdf(html, output)


class ContractGenerator:
    TEMPLATES: ClassVar[dict[str, str]] = {
        "employment": "contracts/employment.html",
        "nda": "contracts/nda.html",
        "termination": "contracts/termination.html",
        "transfer": "contracts/transfer.html",
        "amendment": "contracts/amendment.html",
    }

    def __init__(self, template_dir: str = "templates") -> None:
        self._engine = TemplateEngine(template_dir)
        self._pdf = PDFRenderer()

    def generate(
        self, contract_type: str, data: dict[str, Any], output_format: str = "html"
    ) -> str | bytes:
        template = self.TEMPLATES.get(contract_type)
        if not template:
            raise ValueError(f"Unknown contract type: {contract_type}")

        html = self._engine.render(template, data)

        if output_format == "pdf":
            return self._pdf.html_to_pdf(html)
        elif output_format == "docx":
            return self._html_to_docx(html)
        return html

    def _html_to_docx(self, html: str) -> bytes:
        import io

        from bs4 import BeautifulSoup
        from docx import Document

        soup = BeautifulSoup(html, "html.parser")
        doc = Document()

        for element in soup.body.children if soup.body else []:
            if element.name == "h1":
                doc.add_heading(element.get_text(), level=1)
            elif element.name == "h2":
                doc.add_heading(element.get_text(), level=2)
            elif element.name == "h3":
                doc.add_heading(element.get_text(), level=3)
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name == "ul":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Bullet")
            elif element.name == "table":
                rows = element.find_all("tr")
                if rows:
                    cols = len(rows[0].find_all(["td", "th"]))
                    table = doc.add_table(rows=len(rows), cols=cols)
                    for i, row in enumerate(rows):
                        cells = row.find_all(["td", "th"])
                        for j, cell in enumerate(cells):
                            table.cell(i, j).text = cell.get_text()

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()


class CVGenerator:
    def __init__(self, template_dir: str = "templates") -> None:
        self._engine = TemplateEngine(template_dir)
        self._pdf = PDFRenderer()

    def generate_cv(self, data: dict[str, Any], format: str = "pdf") -> bytes | str:
        html = self._engine.render("cv/modern.html", data)
        if format == "pdf":
            return self._pdf.html_to_pdf(html)
        elif format == "html":
            return html
        raise ValueError(f"Unsupported format: {format}")

    def generate_cover_letter(self, data: dict[str, Any], format: str = "pdf") -> bytes | str:
        html = self._engine.render("cv/cover_letter.html", data)
        if format == "pdf":
            return self._pdf.html_to_pdf(html)
        return html


class LetterGenerator:
    def __init__(self, template_dir: str = "templates") -> None:
        self._engine = TemplateEngine(template_dir)
        self._pdf = PDFRenderer()

    def generate(self, letter_type: str, data: dict[str, Any], format: str = "pdf") -> bytes | str:
        templates = {
            "motivation": "letters/motivation.html",
            "recommendation": "letters/recommendation.html",
            "resignation": "letters/resignation.html",
            "offer": "letters/offer.html",
        }

        template = templates.get(letter_type)
        if not template:
            raise ValueError(f"Unknown letter type: {letter_type}")

        html = self._engine.render(template, data)
        if format == "pdf":
            return self._pdf.html_to_pdf(html)
        return html
