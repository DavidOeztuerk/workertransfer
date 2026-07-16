from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO, ClassVar

import magic
from worker_storage import AzureBlobStorage, LocalStorage, MinIOStorage, S3Storage


@dataclass
class UploadResult:
    key: str
    url: str
    size: int
    content_type: str
    checksum: str


class FileValidator:
    ALLOWED_MIME_TYPES: ClassVar[dict[str, list[str]]] = {
        "image": ["image/jpeg", "image/png", "image/gif", "image/webp", "image/svg+xml"],
        "document": [
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ],
        "spreadsheet": [
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ],
        "video": ["video/mp4", "video/webm", "video/quicktime"],
        "audio": ["audio/mpeg", "audio/wav", "audio/ogg"],
    }

    MAX_FILE_SIZE: ClassVar[int] = 50 * 1024 * 1024  # 50MB

    @classmethod
    def validate(
        cls, file: BinaryIO, allowed_categories: list[str] | None = None
    ) -> tuple[bool, str | None]:
        # Check file size
        file.seek(0, 2)
        size = file.tell()
        file.seek(0)

        if size > cls.MAX_FILE_SIZE:
            return False, f"File size exceeds {cls.MAX_FILE_SIZE} bytes"

        # Check MIME type
        mime = magic.from_buffer(file.read(2048), mime=True)
        file.seek(0)

        allowed = []
        if allowed_categories:
            for cat in allowed_categories:
                allowed.extend(cls.ALLOWED_MIME_TYPES.get(cat, []))
        else:
            for types in cls.ALLOWED_MIME_TYPES.values():
                allowed.extend(types)

        if mime not in allowed:
            return False, f"MIME type {mime} not allowed"

        return True, None


class FileProcessor:
    def __init__(self, storage: S3Storage | MinIOStorage | LocalStorage | AzureBlobStorage):
        self._storage = storage

    async def upload(
        self,
        file: BinaryIO,
        filename: str,
        content_type: str | None = None,
        folder: str = "uploads",
    ) -> UploadResult:
        import hashlib

        # Validate
        valid, error = FileValidator.validate(file)
        if not valid:
            raise ValueError(error)

        # Generate key
        import uuid

        ext = Path(filename).suffix
        key = f"{folder}/{uuid.uuid4()}{ext}"

        # Detect content type if not provided
        if not content_type:
            file.seek(0)
            content_type = magic.from_buffer(file.read(2048), mime=True)
            file.seek(0)

        # Calculate checksum and size (full read).
        file.seek(0)
        content_bytes = file.read()
        file.seek(0)
        checksum = hashlib.sha256(content_bytes).hexdigest()
        size = len(content_bytes)

        # Upload
        url = await self._storage.upload(key, file, content_type)

        return UploadResult(
            key=key,
            url=await self._storage.get_presigned_url(key) if not url else url,
            size=size,
            content_type=content_type,
            checksum=checksum,
        )

    async def upload_multiple(
        self,
        files: list[tuple[BinaryIO, str]],
        folder: str = "uploads",
    ) -> list[UploadResult]:
        results = []
        for file, filename in files:
            results.append(await self.upload(file, filename, folder=folder))
        return results


class ImageProcessor:
    @staticmethod
    async def resize(image: BinaryIO, max_width: int = 1920, max_height: int = 1080) -> BinaryIO:
        import io

        from PIL import Image

        img = Image.open(image)
        img.thumbnail((max_width, max_height))

        output = io.BytesIO()
        img.save(output, format=img.format or "JPEG", quality=85)
        output.seek(0)
        return output

    @staticmethod
    async def convert(image: BinaryIO, format: str = "WEBP", quality: int = 85) -> BinaryIO:
        import io

        from PIL import Image

        img = Image.open(image)
        output = io.BytesIO()
        img.save(output, format=format.upper(), quality=quality)
        output.seek(0)
        return output

    @staticmethod
    async def generate_thumbnails(
        image: BinaryIO, sizes: list[tuple[int, int]] | None = None
    ) -> dict[str, BinaryIO]:
        if sizes is None:
            sizes = [(150, 150), (300, 300), (600, 600)]

        thumbnails = {}
        for width, height in sizes:
            thumb = await ImageProcessor.resize(image, width, height)
            thumbnails[f"{width}x{height}"] = thumb
        return thumbnails


class DocumentProcessor:
    @staticmethod
    async def extract_text(file: BinaryIO, content_type: str) -> str:
        if content_type == "application/pdf":
            import pdfplumber

            with pdfplumber.open(file) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif content_type in [
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]:
            import docx

            doc = docx.Document(file)
            return "\n".join(p.text for p in doc.paragraphs)
        return ""

    @staticmethod
    async def convert_to_pdf(file: BinaryIO, content_type: str, output: BinaryIO) -> None:
        if content_type in [
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]:
            import docx
            from weasyprint import HTML

            doc = docx.Document(file)
            html = "\n".join(f"<p>{p.text}</p>" for p in doc.paragraphs)
            HTML(string=html).write_pdf(output)
        else:
            raise ValueError("Unsupported format for PDF conversion")
